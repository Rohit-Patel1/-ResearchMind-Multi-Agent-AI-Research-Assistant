"""
report_exporter.py
Helpers for exporting the markdown report to PDF and DOCX.
Both functions return bytes so Streamlit can serve them as downloads.
"""
import io
import re


def _strip_markdown(text: str) -> str:
    """Very simple markdown → plain text stripper."""
    text = re.sub(r"#{1,6}\s*", "", text)      
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  
    text = re.sub(r"\*(.+?)\*", r"\1", text)       
    text = re.sub(r"`(.+?)`", r"\1", text)          
    text = re.sub(r"^\s*[-*+]\s+", "• ", text, flags=re.MULTILINE)  
    return text


def export_pdf(markdown_text: str) -> bytes:
    """
    Convert markdown report to PDF bytes.
    Requires: pip install reportlab
    Falls back to fpdf2 if reportlab isn't available.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        )

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=20 * mm,
            rightMargin=20 * mm,
            topMargin=22 * mm,
            bottomMargin=22 * mm,
        )

        styles = getSampleStyleSheet()
        style_body = ParagraphStyle(
            "body",
            parent=styles["Normal"],
            fontSize=10,
            leading=15,
            textColor=colors.HexColor("#333333"),
        )
        style_h1 = ParagraphStyle(
            "h1",
            parent=styles["Heading1"],
            fontSize=18,
            leading=22,
            textColor=colors.HexColor("#111111"),
            spaceAfter=8,
        )
        style_h2 = ParagraphStyle(
            "h2",
            parent=styles["Heading2"],
            fontSize=13,
            leading=17,
            textColor=colors.HexColor("#222222"),
            spaceAfter=5,
            spaceBefore=10,
        )

        story = []
        for line in markdown_text.splitlines():
            stripped = line.strip()
            if stripped.startswith("# "):
                story.append(Paragraph(stripped[2:], style_h1))
            elif stripped.startswith("## "):
                story.append(Paragraph(stripped[3:], style_h2))
            elif stripped.startswith("### "):
                story.append(Paragraph(stripped[4:], style_h2))
            elif stripped == "---":
                story.append(HRFlowable(width="100%", thickness=0.5,
                                        color=colors.HexColor("#cccccc")))
                story.append(Spacer(1, 4))
            elif stripped == "":
                story.append(Spacer(1, 6))
            else:
                # basic inline markdown cleanup for reportlab
                cleaned = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", stripped)
                cleaned = re.sub(r"\*(.+?)\*",     r"<i>\1</i>", cleaned)
                cleaned = re.sub(r"`(.+?)`",        r"\1",        cleaned)
                if stripped.startswith("- ") or stripped.startswith("* "):
                    cleaned = "• " + cleaned[2:]
                story.append(Paragraph(cleaned, style_body))

        doc.build(story)
        return buf.getvalue()

    except ImportError:
        
        try:
            from fpdf import FPDF

            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", size=10)
            plain = _strip_markdown(markdown_text)
            for line in plain.splitlines():
                pdf.multi_cell(0, 6, line)
            return pdf.output(dest="S").encode("latin-1", errors="replace")

        except ImportError:
            raise RuntimeError(
                "PDF export requires 'reportlab' or 'fpdf2'.\n"
                "Install with: pip install reportlab"
            )


def export_docx(markdown_text: str) -> bytes:
    """
    Convert markdown report to DOCX bytes.
    Requires: pip install python-docx
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        import docx.oxml as oxml

        doc = Document()

        
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        for line in markdown_text.splitlines():
            stripped = line.strip()

            if stripped.startswith("# "):
                p = doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                p = doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                p = doc.add_heading(stripped[4:], level=3)
            elif stripped == "":
                doc.add_paragraph("")
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                
                para = doc.add_paragraph()
                parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", stripped)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = para.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith("*") and part.endswith("*"):
                        run = para.add_run(part[1:-1])
                        run.italic = True
                    else:
                        para.add_run(part)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except ImportError:
        raise RuntimeError(
            "DOCX export requires 'python-docx'.\n"
            "Install with: pip install python-docx"
        )